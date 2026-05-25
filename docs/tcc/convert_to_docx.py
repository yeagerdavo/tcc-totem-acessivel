import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Configure basic document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)
        
    # Configure basic document styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(10.0)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            continue
            
        # Handle Headings
        if stripped.startswith('# '):
            title_text = stripped[2:]
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            
            # Align cover elements, resume, abstract, summary, conclusion to center
            center_headings = ['UNIVERSIDADE', 'TOTEM', 'DAVI XAVIER', 'KAUÃ SANTOS', 'RAFAEL LUIZ', 'FOLHA', 'RESUMO', 'SUMÁRIO', 'Santos – SP', '2026']
            if any(term in title_text for term in center_headings):
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = h.add_run(title_text)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
            run.font.size = Pt(14)
            continue
            
        elif stripped.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[3:])
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
            run.font.size = Pt(13)
            continue
            
        elif stripped.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[4:])
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            run.bold = True
            run.font.size = Pt(12)
            continue
            
        elif stripped.startswith('#### '):
            h = doc.add_heading(level=4)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(stripped[5:])
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            run.bold = True
            run.font.size = Pt(11)
            continue
            
        # Handle empty lines
        if not stripped:
            continue
            
        # Handle bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
        else:
            content = line.rstrip('\n')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
        # Parser for bold **text** and italics *text*
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)
                
    doc.save(docx_path)
    print("Conversion completed successfully!")

if __name__ == '__main__':
    md_file = r'c:\Users\user\Documents\PROJETOS\TCC\tcc-totem-acessivel\docs\tcc\documento_tcc_totem.md'
    docx_file = r'c:\Users\user\Documents\PROJETOS\TCC\tcc-totem-acessivel\docs\tcc\documento_tcc_totem.docx'
    convert_md_to_docx(md_file, docx_file)
